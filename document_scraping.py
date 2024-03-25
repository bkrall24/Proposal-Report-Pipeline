import numpy as np
import re
from docx import Document, table, text
import pandas as pd

def document_TOC(input, toc_name = "TABLE OF CONTENTS"):
    
    toc_ind = [ind for ind,p in enumerate(input.paragraphs) if toc_name.lower() in p.text.lower()]
    # toc_ind = [ind for ind,p in enumerate(input.paragraphs) if p.text == toc_name]
    toc = [p.text for p in input.paragraphs[toc_ind[0]:] if bool(re.search(r'\t+\d', p.text))]
    # toc = []
    # c_ind= []
    # for ind, p in enumerate(input.paragraphs):
    #     if bool(re.search(r'\t+\d', p.text)):
    #         toc.append(p.text)
    #         c_ind.append(ind)

    titles = []
    levels = []
    for line in toc:
        split_line = line.split('\t')
        if len(split_line) == 3:
            sec = split_line[0]
            nums = sec.split('.')
            titles.append(split_line[-2].strip())
            levels.append(len(nums))

    index = np.array([ind for ind, paragraph in enumerate(input.paragraphs) if paragraph.text.strip() in titles and "Heading" in paragraph.style.name])
    

    if len(index) != len(titles):
        found = [titles.index(input.paragraphs[i].text.strip()) for i in index]
        titles = [titles[i] for i in found]
        levels = [levels[i] for i in found]

    titles = [t.strip().replace(' ', '_') for t in titles]
    data_ind = {}
    data_ind["Title_Page"] = (0, toc_ind[0])
    data_ind["TABLE_OF_CONTENTS"] = (toc_ind[0], index[0])

    for ind, title in enumerate(titles[:-1]):

        start = index[ind]
        next_ind = [i for i,element in enumerate(levels) if element <= levels[ind] and i > ind]
        if len(next_ind):
            end = index[next_ind][0]
        else:
            end = len(input.paragraphs)
        data_ind[title] = (start, end)

    data_ind[titles[-1]] = (index[-1], len(input.paragraphs))

    return data_ind


def document_tables(doc):

    para_count = 0
    table_loc = []
    tables = []
    for p in doc.iter_inner_content():
       
        if isinstance(p, table.Table):
            table_loc.append(para_count)
            tables.append(p)
        elif isinstance(p, text.paragraph.Paragraph):
            para_count += 1
    
    return table_loc, tables


def scrape_proposal_data(input_path, name_doc= "/Users/rebeccakrall/Data/Proposal-Report-Data/Reference_Data/Melior_people.xlsx",
                         assay_doc = "/Users/rebeccakrall/Data/Proposal-Report-Data/Reference_Data/Assay_Codes.xlsx"):

    input = Document(input_path)

    # Pull data from proposal (input) 
    data_ind = document_TOC(input)    
    table_loc, table = document_tables(input)
    data = {}

    d1 = scrape_project_data(0, data_ind['Title_Page'][1], input)
    # title_page = [p.text.strip() for p in input.paragraphs[0:data_ind['Title_Page'][1]] if p.text.strip() != ''] 

    # # project info
    # data['title'] = title_page[0]
    # company_line = [x.split('Prepared for ')[-1] for x in title_page if 'Prepared for' in x]
    # if len(company_line) >0:
    #     data['company'] = company_line[0]

    # desc = input.sections[0].header.paragraphs[0].text.split('\t')
    # if len(desc):
    #     data['description'] = desc[0]

    # sn = input.sections[0].header.paragraphs[1].text.split('\t')
    # if len(sn) > 0:
    #     if ':' in sn:
    #         data['study_num'] = sn.split(':')[-1].strip()[0]
    #     else:
    #         data['study_num'] = sn[0]

    #     data['company_acronym'] = data['study_num'].split('_')[0].strip()
    #     data['method_abbreviation'] = data['study_num'].split('_')[3].strip()

    # dates = [x.split(':')[-1].strip() for x in title_page if 'Date' in x]
    # if len(dates) > 0:
    #     data['issue'] = dates[0]

    # if len(dates) > 1:
    #     data['reissues'] = dates[1:]
   
    # experimental information
    k = find_string_list(['Animal_Description'], data_ind.keys())
    if k is not None:
        start,end = data_ind[k]
        ad_dict = delineated_para_list(input.paragraphs[start:end])

        for k in ad_dict.keys():
            if 'species' in k.lower():
                data['species'] = ad_dict[k]
            
            if 'strain' in k.lower():
                data['strain'] = ad_dict[k]
            
            if 'age' in k.lower():
                data['age'] = ad_dict[k]

    k = find_string_list(['Design'], data_ind.keys())
    if k is not None:
        start,end = data_ind[k]
        ad_dict = delineated_para_list(input.paragraphs[start:end])

        for k in ad_dict.keys():
            if 'number of groups' in k.lower():
                data['group'] = ad_dict[k]
            
            if 'number of animals per group' in k.lower():
                data['n'] = ad_dict[k]
            

    # People
    names = pd.read_excel(name_doc)
    names = names['Names'].values

    k = find_string_list(['Project_Manager'], data_ind.keys())
    if k is not None:
        start,end = data_ind[k]
        pm = find_string_paragraphs(input.paragraphs[start:end], names)
        if pm is None:
            pm = input.paragraphs[start+1].text.split("will")[0]
        data['project_manager'] = pm
    
    k = find_string_list(['Project_Coordinator'], data_ind.keys())
    if k is not None:
        start,end = data_ind[k]
        pc = find_string_paragraphs(input.paragraphs[start:end], names)
        if pc is None:
            pc = input.paragraphs[start+1].text.split("will")[0]
        data['project_coordinator'] = pc

    k = find_string_list(['Principal_Associate'], data_ind.keys())
    if k is not None:
        start,end = data_ind[k]
        pa = find_string_paragraphs(input.paragraphs[start:end], names)
        if pa is None:
            pa = input.paragraphs[start+1].text.split("will")[0]
        data['principal_associate'] = pa

    k = find_string_list(['Client_Management_Specialist'], data_ind.keys())
    if k is not None:
        start,end = data_ind[k]
        cms = find_string_paragraphs(input.paragraphs[start:end], names)
        if cms is None:
            cms = input.paragraphs[start+1].text.split("will")[0]
        data['client_management_specialist'] = cms

    # k = find_string_list(['Methods'], data_ind.keys())
    # assays = pd.read_excel(assay_doc)
    # ass = assays['Assay']
    # k = find_string_list(['Appendix'] ,data_ind.keys())
    # if k is not None:
    #     end = data_ind[k][1]
    # else:
    #     end = len(input.paragraphs)
    #     data['methods'] = [p.text for p in input.paragraphs[start+1:end] if "Heading" in p.style.name]
    
    methods, codes, method_headings = scrape_method_data(input, data_ind)
    data['methods'] = methods
    data['codes'] =codes
    data['methods_headings'] = method_headings
 
    data = {**data, **d1}

    # Abbreviations
    return data, data_ind


def delineated_para_list(doc, delineator = ':'):
    dat = {}
    last_key = None
    for p in doc:
        split = p.text.split(delineator, 1)
        if len(split) > 1 and not check_if_time(split[-1]):
            split = p.text.split(delineator, 1)
            last_key = split[0]
            dat[last_key] = split[-1].strip()
        elif not last_key is None:
            old_val = dat[last_key]

            if old_val != '':
                dat[last_key] = old_val + ", " + p.text.strip()
            else:
                dat[last_key] = p.text.strip()
    
    return dat



def check_if_time(txt):
    found = False

    ind = txt.find('AM')
    if (ind < 5) & (ind > -1):
        found = True
    
    ind = txt.find('PM')
    if (ind < 5) & (ind > -1):
        found = True
    
    return found
    


def find_string_paragraphs(doc, str_list, all= False):
    found_names = []
    for p in doc:
        for s in str_list:
            
            if s.lower() in p.text.lower():
                if all:
                    found_names.append(s)
                else:
                    return s
                # print(s)
    
    # return found_names
    return found_names

def find_string_list(str_list1, str_list2):
    for p in str_list1:
        for s in str_list2:
            if s.lower() in p.lower():
                return s
            elif p.lower() in s.lower():
                return s
    
    return None


def build_database(file_list):
    all_dicts = []
    working_files = []
    not_working_files = []
    for f in file_list:
        try:
            d, di = scrape_proposal_data(f)
            d['Pathfile'] = f
            all_dicts.append(d)
            working_files.append(f)
        except Exception as e:
            print("Error: ",e)
            not_working_files.append(f)

    df = pd.DataFrame(all_dicts)
    df.to_csv("/Users/rebeccakrall/Desktop/Proposal_data2_2023.csv")
    
    return all_dicts, working_files, not_working_files

def collect_abbreviations(file_list):
    all_dicts = []
    for f in file_list:
        try:
            abbs = scrape_abbreviations(f)
            all_dicts.append(abbs)
        except Exception as e:
            print("Error: ",e)
        
    
    df = pd.DataFrame(all_dicts)
    df.to_csv("/Users/rebeccakrall/Desktop/Proposal_data2_2023.csv")

def scrape_project_data(start, end, input):

    data = {}
    title_page = [p.text.strip() for p in input.paragraphs[start:end] if p.text.strip() != '']
    data['title'] = title_page[0]


    company_line = [x.split('Prepared for ')[-1] for x in title_page if 'Prepared for' in x]
    if len(company_line) >0:
        data['client'] = company_line[0]

    desc = input.sections[0].header.paragraphs[0].text.split('\t')
    if len(desc):
        data['description'] = desc[0]

    sn = input.sections[0].header.paragraphs[1].text.split('\t')
    if len(sn) > 0:
        if ':' in sn[0]:
            data['study_num'] = sn[0].split(':')[-1].strip()
        else:
            data['study_num'] = sn[0]

        # print(data['study_num'])
        data['company_acronym'] = data['study_num'].split('_')[0].strip()
        data['method_abbreviation'] = data['study_num'].split('_')[3].strip()

    dates = [x.split(':')[-1].strip() for x in title_page if 'Date' in x]
    if len(dates) > 0:
        data['issue'] = dates[0]

    if len(dates) > 1:
        data['reissues'] = dates[1:]

    return data


def scrape_people_data(data_ind, input, name_doc= "/Users/rebeccakrall/Data/Proposal-Report-Data/Reference_Data/Melior_people.xlsx" ):
    names = pd.read_excel(name_doc)
    names = names['Names'].values

    data = {}
    k = find_string_list(['Project_Manager'], data_ind.keys())
    if k is not None:
        start,end = data_ind[k]
        pm = find_string_paragraphs(input.paragraphs[start:end], names)
        if pm is None:
            pm = input.paragraphs[start+1].text.split("will")[0]
        data['pm'] = pm
    
    k = find_string_list(['Project_Coordinator'], data_ind.keys())
    if k is not None:
        start,end = data_ind[k]
        pc = find_string_paragraphs(input.paragraphs[start:end], names)
        if pc is None:
            pc = input.paragraphs[start+1].text.split("will")[0]
        data['pc'] = pc

    k = find_string_list(['Principal_Associate'], data_ind.keys())
    if k is not None:
        start,end = data_ind[k]
        pa = find_string_paragraphs(input.paragraphs[start:end], names)
        if pa is None:
            pa = input.paragraphs[start+1].text.split("will")[0]
        data['pa'] = pa

    k = find_string_list(['Client_Management_Specialist'], data_ind.keys())
    if k is not None:
        start,end = data_ind[k]
        cms = find_string_paragraphs(input.paragraphs[start:end], names)
        if cms is None:
            cms = input.paragraphs[start+1].text.split("will")[0]
        data['cms'] = cms

    return data

def scrape_method_data(doc, data_ind, assay_doc = "/Users/rebeccakrall/Data/Proposal-Report-Data/Reference_Data/Assay_List.xlsx"):
    assays = pd.read_excel(assay_doc)
    ass = assays['Assay']
    k = find_string_list(['Experimental_Procedures'] ,data_ind.keys())
    if k is not None:
        start = data_ind[k][0]
        end = data_ind[k][1]
    else:
        start = 0
        end = len(doc.paragraphs)

    # methods = [p.text for p in doc.paragraphs[start+1:end] if "Heading" in p.style.name]
    methods = list(set(find_string_paragraphs(doc.paragraphs[start:end], ass, all = True)))

    # print(set(methods))
    codes = []
    for s in methods:
        codes.append(assays['Code'][assays['Assay'] == s].values[0])


    k = find_string_list(['Methods'] ,data_ind.keys())
    if k is not None:
        start = data_ind[k][0]
        end = data_ind[k][1]
        method_headings = [p.text for p in doc.paragraphs[start+1:end] if "Heading" in p.style.name]
    else:
        method_headings = []
    

    return methods, codes, method_headings

def scrape_abbreviations(doc, data_ind):
    k = find_string_list(['Abbreviations'], data_ind.keys())
    if k is not None:
        start,end = data_ind[k]
        abb_dict = delineated_para_list(doc.paragraphs[start:end])

        return abb_dict
    else:
        return None
    
def scrape_design_data(input, data_ind):

    d = {}
    k = find_string_list(['Design'], data_ind.keys())
    if k is not None:
        start,end = data_ind[k]
        ad_dict = delineated_para_list(input.paragraphs[start:end])

        standard_design = ["dose volume", "formulation", "dose levels", "dose frequency", "study duration", "pretreatment time", "number of groups",
                        "number of animals per group", "total number of animals"]

        
        for k in ad_dict.keys():
            found = False
            for s in standard_design:
                if (k.lower() in s) or (s in k.lower()):
                    d[s] = ad_dict[k]
                    found = True

            if ("route(s) of administration" in k.lower()) or (("route" in k.lower()) and ("adminstration" in k.lower())):
                d["route of administration"] = ad_dict[k]
                found = True
            
            if found is False:
                d[k] = ad_dict[k]
             
    else:
        d = None
    return d



def scrape_design_tables(tables, study_id):
    all_design = []
    for t in tables:
        try:
            mdf = get_design_table(t, study_id)
            if mdf is not None:
                all_design.append(mdf)
        except:
            print('table failed')
    
    merged_df = pd.concat(all_design, ignore_index = True)
    return merged_df

def get_design_table(table, study_id):
    
    # assume the first row contains table headers
    cols = [cell.text.lower() for cell in table.rows[0].cells]

    # make a dataframe containing the remaining rows
    data = []
    for i, row in enumerate(table.rows[1:]):
        text = [cell.text.lower() for cell in row.cells]

        row_data = dict(zip(cols, text))
        data.append(row_data)
    df = pd.DataFrame(data)

    # identify columns containing 'treatment' to add as option in table
    treatments = [a for  a in df.columns if 'treatment' in a]

    # identify expected columns
    col_ind = []
    rn = {}
    target_cols = ['group #', 'group size', 'days', 'dose', 'route']
    for ind, c in enumerate(cols):
        for t in target_cols:
            if t in c:
                col_ind.append(c)
                rn[c] = t
    
    # pull out corresponding dataframes for each treatment, adding study_id 
    df['study_id'] = study_id
    all_treats = []
    for t in treatments:
        rn[t] = 'treatment'
        treat = df.loc[:, ['study_id'] + [t] + col_ind]
        treat = treat.rename(columns = rn)
        all_treats.append(treat)

    # combine to a single dataframe
    if len(all_treats):
        merged_df = pd.concat(all_treats, ignore_index = True)
    else:
        merged_df = None
    

    return merged_df