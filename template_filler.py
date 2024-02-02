from docxtpl import DocxTemplate, subdoc
from docx import Document, table, text
import re
import numpy as np
from datetime import date, datetime
import os
import pyinflect
import spacy
import re
from docx.enum.text import WD_COLOR_INDEX
# import pandas as pd


def find_subject_plural(token):
    # subjs = [t for t in token.ancestors if 'subj' in t.dep_]
    subjs = [t for t in token.lefts if 'subj' in t.dep_]
    # nouns = [t for t in token.lefts if 'NOUN' in t.dep_]
    nouns = [t for t in token.sent[:token.i] if 'NOUN' in t.pos_]
    if subjs:

        # print(subjs[0])
        morph = subjs[0].morph.to_dict()
        if 'Number' in morph.keys():
            return morph['Number']
        elif 'S' in subjs[0].tag_:
            return 'Plur'
        else:
            return 'Sing'
            
            
    elif nouns:
        # print(nouns[0])
        morph = nouns[-1].morph.to_dict()
        if 'Number' in morph.keys():
            return morph['Number']
        elif 'S' in nouns[-1].tag_:
            return 'Plur'
        else:
            return 'Sing'
    
    else:
        return 'Sing'

def past_tensify(txt, nlp):
    doc = nlp(txt)
    out = list()
    words = []
    aux_count = 0
    aux_phrase = []

    for word in doc:
        words.append(word)
        new_word = None
        tag = word.tag_
        pos = word.pos_
        dep = word.dep_

        if 'AUX' in pos and dep != 'ROOT' and out[-1].text.lower() != 'to':# and word.text != 'is':
            aux_count = aux_count +1
            aux_phrase.append(word.lemma_)

            # remove auxillary words, but keep track of recent removals. 
            # Do not remove auxillary words that are roots (i.e. be as the root verb)
            # Do not remove auxillary words following 'to'
            
        else:  
            if (dep == "ROOT") or ('cl' in dep) or ('conj' in dep) :

                # if dep == 'ROOT':
                #     print("ROOT: "+word.text)
                
                # if 'cl' in dep:
                #     print("cl: "+word.text)

                
                # determine if the subject is plural or singular
                plural = find_subject_plural(word)

                # pull out any adverbs
                adverbs = [a for a in word.children if 'RB' in a.tag_]

                # Identify 'be' in auxillary clause (will be, have been) to replace with was/were
                # or Identify 'be' in conjuction with past participle (are having, am walking) to replace with was/were
                if ((('be' in aux_phrase) or ('am' in aux_phrase) ) and aux_count > 1) or \
                    ((('be' in aux_phrase) or ('am' in aux_phrase) ) and word.tag_ in ['VBN', 'VBG']): 
                    # print('AUX lemmas \t')
                    # print(aux_phrase)

                    if word.tag_ != 'VBG':
                        if plural == 'Sing':
                            out.append(nlp('was')[0])
                        else:
                            out.append(nlp('were')[0])
                    else:
                        new_word = word._.inflect('VBD', form_num = 0)
                
                    aux_count = 0
                    aux_phrase = []

                if (tag in ['VB', 'VBP']) and ( word.text.lower() != 'see') and out[-1].text.lower() != 'to': # removed VBN and VBZ in wild test # 

                    if word.lemma_ == 'be':
                        if plural == 'Sing':
                            new_word = 'was'
                        else:
                            new_word = 'were'
                    else:      
                        if plural == 'Sing':
                            new_word = word._.inflect('VBD', form_num = 0)

                            # print('Inflection: '+ new_word)
                        else:
                            new_word = word._.inflect('VBD', form_num = 1)
                            # print('Inflection: '+ new_word)

                elif new_word is None:
                    new_word = word.text
                
                for adv in adverbs:
                    # print("ADVERB: "+adv.text)
                    if adv in out and adv.i in range(word.i - (len(adverbs)+1), word.i + (len(adverbs)+1)):
                        
                        out.pop(out.index(adv))
                        out.append(adv)

                if new_word is not None:
                    out.append(nlp(new_word)[0])
                else:
                    out.append(word)
                
                aux_count = 0
                aux_phrase = []
            
            else:
                out.append(word)

    
            
    all_words = [" "+t.text if t.pos_ != 'PUNCT' else t.text for t in out]

    # print(out)
    if len(out) > 0:
        out = "".join(all_words)

    return out     
    
  

def tense_correct_para(txt):
    nlp = spacy.load('en_core_web_sm')
    delimiters = r'[.:;\t]'
    chunks = re.split(delimiters, txt)
    for c in chunks:
        new_chunk = past_tensify(c, nlp)
        # print(c)
        if len(new_chunk) > 0:
            # print(new_chunk)
            txt = txt.replace(c, new_chunk)

    return txt

# def tense_correct_style_match(para):
#     nlp = spacy.load('en_core_web_sm')

#     full_style = para.style
#     for c in para.runs:
#         italics = c.italic
#         bold = c.bold
#         underline = c.underline
#         fontname = c.font.name
#         style = c.style
    
#     delimiters = r'[.:;\t]'
#     chunks = re.split(delimiters, para.text)
#     for c in chunks:
#         new_chunk = past_tensify(c, nlp)
#         # print(c)
#         if len(new_chunk) > 0:
#             para.text = para.text.replace(c, new_chunk)
    
#     para.style = full_style
#     for c in para.runs:
#         c.italics = italics
#         c.bold = bold 
#         c.underline = underline 
#         c.font.name = fontname 
#         c.style = style 
    
#     return para 
        
        



# def font_matching(run):
    
#     new_run.italic = old_run.italic
#     new_run.bold = old_run.bold
#     new_run.underline = old_run.underline
#     new_run.font = old_run.font
#     new_run.style = old_run.style
    


def parse_toc(doc):
    # potentially want to add functionality to figure out subheadings - dict in dict, json?, 
    toc_ind = [ind for ind,p in enumerate(doc.paragraphs) if p.text == 'TABLE OF CONTENTS']
    toc = [p.text for p in doc.paragraphs[toc_ind[0]:] if bool(re.search(r'\t+\d', p.text))]

    titles = []
    page_ind = []
    for line in toc:
        split_line = line.split('\t')
        # line_ind.append(float(split_line[0]))
        titles.append(split_line[-2].strip())
        page_ind.append(int(split_line[-1]))


    index = np.array([ind for ind, paragraph in enumerate(doc.paragraphs) if paragraph.text.strip() in titles])

    data = {}
    data_ind = {}
    data_ind["Title Page"] = 0
    data_ind["TABLE OF CONTENTS"] = toc_ind[0]
    for start,end in zip(index, np.append(index[1:], len(doc.paragraphs))):
        key = doc.paragraphs[start].text.strip()
        key = key.replace(' ', '_')

        data[key] = " ".join([p.text for p in doc.paragraphs[start+1:end]])
        data_ind[key] = start
    
    sortem = sorted(data_ind.items(), key=lambda x: x[1]) 
    data_ind = {tup[0]:tup[1] for tup in sortem}
    return data_ind , data

def find_tables(doc):
    para_count = 0
    table_loc = []
    tables = []
    for p in doc.iter_inner_content():
       
        if isinstance(p, table.Table):
            table_loc.append(para_count)
            tables.append(p)
        elif isinstance(p, text.paragraph.Paragraph):
            para_count += 1
    
    return table_loc, tables




def move_table_after(table, paragraph):
    tbl, p = table._tbl, paragraph._p
    p.addnext(tbl)


def scrape_proposal(input_path, template_path, save_path = None, save = True):
    
    # Load proposal and template docx
    input = Document(input_path)
    template = DocxTemplate(template_path)

    # Pull data from proposal (input)
    data_ind,data = parse_toc(input) 
    title_page =[(ind, p.text) for ind, p in enumerate(input.paragraphs) if ind < data_ind['TABLE OF CONTENTS'] and p.text.strip() != '' ]    
    table_loc, table = find_tables(input)

    # Pull placeholders from template
    place = template.get_undeclared_template_variables()
    replace = {k : None for k in place}

    # Identify key points, update replace dict
    replace['title'] = title_page[0][1]
    replace['study_num'] = [ x[1].split(': ')[-1] for x in title_page if 'Project Quotation' in x[1]][0]
    replace['client'] = [x[1].split('Prepared for ')[-1] for x in title_page if 'Prepared for' in x[1]][0]
    replace['t'] = date.today().strftime("%B %d, %Y")
    replace['description'] = input.sections[0].header.paragraphs[0].text.split('\t')[0]

    # Roughly identify the project manager and coordinator
    m = [x for x in data.keys() if 'Project_Manager' in x]
    if len(m) > 0:
        replace['pm'] = " ".join(data[m[0]].split('proposes ')[-1].split(" ")[0:2])

    c = [x for x in data.keys() if 'Project_Coordinator' in x]
    if len(c) > 0:
        replace['pc'] = " ".join(data[c[0]].split('proposes ')[-1].split(" ")[0:2])

    # Replace '&' in strings to prevent jinga2 errors
    for k,v in replace.items():
        if v is not None and '&' in v:
            replace[k] = v.replace("&","and")

    
    # For section replacement, iterate through and pull paragraphs from proposal
    all_ind = np.array(list(data_ind.values()))
    for p in place:
        if p in data_ind.keys():
            
            start = data_ind[p]
            end = all_ind[np.argwhere(all_ind == start)[0][0]+1]
            # print(p+": "+str(start)+":"+str(end))
            sd1 = template.new_subdoc()
            
          
            for x in range(start+1, end, 1):
                if x in table_loc:
                    move_table_after(table[table_loc.index(x)], sd1.paragraphs[-1])
                
                if p != 'References':
                    # para = highlight_TBDs(input.paragraphs[x])
                    # new_text = tense_correct_para(input.paragraphs[x].text)
                    new_text = tense_correct_para(input.paragraphs[x].text)
                    sd1.add_paragraph(new_text, input.paragraphs[x].style)
                    # sd1.add_paragraph(new_para.text, new_para.style)
                else:
                    sd1.add_paragraph(input.paragraphs[x].text, input.paragraphs[x].style)

            # print(start)
            # print(end)
            sd1 = style_match(input.paragraphs[int(start+1) :int(end)], sd1)
            replace[p] = highlight_TBDs(sd1)
    
    path, fn = os.path.split(input_path)
    name, ext = os.path.splitext(fn)
    new_fn = name + '_REPORT' + ext


    if save_path is None:
        save_path = os.path.join(path, new_fn)
    else:
        save_path = os.path.join(save_path, new_fn)


    template.render(replace)

    if save:
        template.save(save_path)

    return template, save_path



def highlight_TBDs(sd1):
    
    to_highlight = ['TBD','tbd', 'to be determined', 'To be determined']
    # print(to_highlight)
    
    if isinstance(sd1, table.Table):
        for h in to_highlight:
            for x in sd1._cells:
                for p in x.paragraphs:
                    for run in p.runs:
                        if h in run.text:
                            run.font.highlight_color= WD_COLOR_INDEX.YELLOW

    elif isinstance(sd1, subdoc.Subdoc):
        for h in to_highlight:
            for para in sd1.paragraphs:
                if h in para.text:
                    for run in para.runs:
                        if h in run.text:
                            run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                                
    return sd1

# def style_match(oldsubdoc, newsubdoc):

#     for oldpara, newpara in zip(oldsubdoc, newsubdoc.paragraphs):
#         newpara.style = oldpara.style
#         for oldrun, newrun in zip(oldpara.runs, newpara.runs):
        
#             newrun.italic = oldrun.italic
#             newrun.bold = oldrun.bold
#             newrun.underline = oldrun.underline
#             newrun.font.name = oldrun.font.name
#             newrun.style = oldrun.style
           
                            
#     return newsubdoc
       


if __name__ == "__main__":

    template_path = "/Users/rebeccakrall/Desktop/Example for Proposal Report Automation/Report_Template_dxpt.docx"
    proposal_dir = "/Volumes/O-Drive/Clients/_PROPOSALS FOR REPORT TEMPLATE_Becca/Proposals"
    report_dir = "/Volumes/O-Drive/Clients/_PROPOSALS FOR REPORT TEMPLATE_Becca/Reports"
    # proposal_dir = "/Users/rebeccakrall/Desktop/Example for Proposal Report Automation/Test_Proposal"
    # report_dir = "/Users/rebeccakrall/Desktop/Example for Proposal Report Automation/Report"

    current_proposals = os.listdir(proposal_dir)
    current_reports = os.listdir(report_dir)

    proposal_names = [x.split('.docx')[0] for x in current_proposals]
    report_names = [x.split('_REPORT.docx')[0] for x in current_reports]

    unmatched = [p for r,p in zip(proposal_names, current_proposals) if r not in report_names]
    for u in unmatched:
        proposal_path = os.path.join(proposal_dir, u)
        template, report_path =scrape_proposal(proposal_path, template_path, save_path = report_dir)



    logfile = "/Users/rebeccakrall/Code/Proposal-Report-Pipeline/report.txt"
    with open(logfile, 'a') as f:
        if len(unmatched) > 0:
            t = datetime.now().strftime("%m/%d/%Y, %H:%M:%S")
            f.write(t)
            f.write('\n')
            for u in unmatched:
                f.write(u + " Report generated")
                f.write('\n')
    
    # print("All proposals scraped")
    