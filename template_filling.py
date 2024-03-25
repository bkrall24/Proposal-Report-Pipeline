from docxtpl import DocxTemplate, subdoc, InlineImage
from docx import Document, table, text
from docx.shared import Inches, Mm

import re
import numpy as np
from datetime import date, datetime
import os
# 
import spacy
# import pyinflect
# import en_core_web_sm
# import lemminflect
import re
from docx.enum.text import WD_COLOR_INDEX
from docx.enum.style import WD_STYLE_TYPE
from document_scraping import *
from docx2python import docx2python
from past_tense import *


def extract_tabs(filepath):
    tab_dict = {}
    input = Document(filepath)
    # extract docx content
    with docx2python(filepath) as docx_result:
        for ind, p in enumerate(input.paragraphs):
            # if 'Bullet' in p.style.name:
            loc = docx_result.text.find(p.text)
            last_nl = docx_result.text.rfind('\n', 0, loc)
            c = docx_result.text[last_nl:loc].count('\t')
            if c > 0:
                tab_dict[ind] = docx_result.text[last_nl:loc].count('\t')
        
    return tab_dict

def list_number(doc, par, prev=None, level=None, num=True):
    """
    Makes a paragraph into a list item with a specific level and
    optional restart.

    An attempt will be made to retreive an abstract numbering style that
    corresponds to the style of the paragraph. If that is not possible,
    the default numbering or bullet style will be used based on the
    ``num`` parameter.

    Parameters
    ----------
    doc : docx.document.Document
        The document to add the list into.
    par : docx.paragraph.Paragraph
        The paragraph to turn into a list item.
    prev : docx.paragraph.Paragraph or None
        The previous paragraph in the list. If specified, the numbering
        and styles will be taken as a continuation of this paragraph.
        If omitted, a new numbering scheme will be started.
    level : int or None
        The level of the paragraph within the outline. If ``prev`` is
        set, defaults to the same level as in ``prev``. Otherwise,
        defaults to zero.
    num : bool
        If ``prev`` is :py:obj:`None` and the style of the paragraph
        does not correspond to an existing numbering style, this will
        determine wether or not the list will be numbered or bulleted.
        The result is not guaranteed, but is fairly safe for most Word
        templates.
    """
    xpath_options = {
        True: {'single': 'count(w:lvl)=1 and ', 'level': 0},
        False: {'single': '', 'level': level},
    }

    def style_xpath(prefer_single=True):
        """
        The style comes from the outer-scope variable ``par.style.name``.
        """
        style = par.style.style_id
        return (
            'w:abstractNum['
                '{single}w:lvl[@w:ilvl="{level}"]/w:pStyle[@w:val="{style}"]'
            ']/@w:abstractNumId'
        ).format(style=style, **xpath_options[prefer_single])

    def type_xpath(prefer_single=True):
        """
        The type is from the outer-scope variable ``num``.
        """
        type = 'decimal' if num else 'bullet'
        return (
            'w:abstractNum['
                '{single}w:lvl[@w:ilvl="{level}"]/w:numFmt[@w:val="{type}"]'
            ']/@w:abstractNumId'
        ).format(type=type, **xpath_options[prefer_single])

    def get_abstract_id():
        """
        Select as follows:

            1. Match single-level by style (get min ID)
            2. Match exact style and level (get min ID)
            3. Match single-level decimal/bullet types (get min ID)
            4. Match decimal/bullet in requested level (get min ID)
            3. 0
        """
        for fn in (style_xpath, type_xpath):
            for prefer_single in (True, False):
                xpath = fn(prefer_single)
                ids = numbering.xpath(xpath)
                if ids:
                    return min(int(x) for x in ids)
        return 0

    if (prev is None or
            prev._p.pPr is None or
            prev._p.pPr.numPr is None or
            prev._p.pPr.numPr.numId is None):
        if level is None:
            level = 0
        numbering = doc.part.numbering_part.numbering_definitions._numbering
        # Compute the abstract ID first by style, then by num
        anum = get_abstract_id()
        # Set the concrete numbering based on the abstract numbering ID
        num = numbering.add_num(anum)
        # Make sure to override the abstract continuation property
        num.add_lvlOverride(ilvl=level).add_startOverride(1)
        # Extract the newly-allocated concrete numbering ID
        num = num.numId
    else:
        if level is None:
            level = prev._p.pPr.numPr.ilvl.val
        # Get the previous concrete numbering ID
        num = prev._p.pPr.numPr.numId.val
    par._p.get_or_add_pPr().get_or_add_numPr().get_or_add_numId().val = num
    par._p.get_or_add_pPr().get_or_add_numPr().get_or_add_ilvl().val = level  

def highlight_TBDs(sd1):
    
    to_highlight = ['TBD','tbd', 'to be determined', 'To be determined']
    # print(to_highlight)
    
    if isinstance(sd1, table.Table):
        for h in to_highlight:
            for x in sd1._cells:
                for p in x.paragraphs:
                    for run in p.runs:
                        if h in run.text:
                            run.font.highlight_color= WD_COLOR_INDEX.YELLOW

    elif isinstance(sd1, subdoc.Subdoc):
        for h in to_highlight:
            for para in sd1.paragraphs:
                if h in para.text:
                    for run in para.runs:
                        if h in run.text:
                            run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                                
    return sd1

def update_table_text(sd1):
    if isinstance(sd1, table.Table):
        for x in sd1._cells:
            for p in x.paragraphs:
                new_text = tense_correct_para(p.text)
                # print("UPDATED TABLE TEXT")
                if new_text != p.text:
                    p.text = new_text
    
    return sd1

def move_table_after(table, paragraph):
    tbl, p = table._tbl, paragraph._p
    p.addnext(tbl)

def style_match(oldsubdoc, newsubdoc):

    for oldpara, newpara in zip(oldsubdoc, newsubdoc.paragraphs):
        newpara.style = oldpara.style
        for oldrun, newrun in zip(oldpara.runs, newpara.runs):
        
            newrun.italic = oldrun.italic
            newrun.bold = oldrun.bold
            newrun.underline = oldrun.underline
            newrun.font.name = oldrun.font.name
            newrun.style = oldrun.style
           
                            
    return newsubdoc

def find_signature(name, sig_dir = "/Users/rebeccakrall/Data/Proposal-Report-Data/Reference_Data/Signatures"):
    
    filename = find_string_list([name], os.listdir(sig_dir))
    if not filename is None:
        full_path = os.path.join(sig_dir, filename)
        return full_path
    else:
        return None


def proposal_to_report(input_path, save_path = None, save = True, template_path = "/Users/rebeccakrall/Data/Proposal-Report-Data/Templates/Report_Template_dxpt.docx"):

    input = Document(input_path)
    template = DocxTemplate(template_path)
    # tab_dict = extract_tabs(input_path)

    # get data of interest
    data_ind = document_TOC(input)
    table_loc, table = document_tables(input)
    proj = scrape_project_data(0, data_ind['Title_Page'][1], input)
    people = scrape_people_data(data_ind, input)
   

    place = template.get_undeclared_template_variables()
    replace = {k : None for k in place}

    for k in replace:
        if k in proj.keys():
            replace[k] = proj[k]
        if k in people.keys():
            replace[k] = people[k]
    
    replace['t'] = date.today().strftime("%B %d, %Y")

    for k,v in replace.items():
        if v is not None and '&' in v:
            replace[k] = v.replace("&","and")

    for p in place:
        if p in data_ind.keys():
            
            start = data_ind[p][0]
            end = data_ind[p][1]
            sd1 = template.new_subdoc()

            for x in range(start+1, end, 1):
                if x in table_loc:
                    updated_table = update_table_text(highlight_TBDs(table[table_loc.index(x)]))

                    move_table_after(updated_table, sd1.paragraphs[-1])

                if p != 'References':
                    if 'Bullet' in input.paragraphs[x].style.name:# or 'List' in input.paragraphs[x].style.name:
                        new_text = tense_correct_para(input.paragraphs[x].text)
                        para = sd1.add_paragraph(new_text, input.paragraphs[x].style)

                        # if x in tab_dict.keys():
                        #     lev = tab_dict[x]
                        # else:
                        lev = None
                        list_number(sd1, para, prev=input.paragraphs[x], level=lev, num=False)
                    else:
                        new_text = tense_correct_para(input.paragraphs[x].text)
                        sd1.add_paragraph(new_text, input.paragraphs[x].style)
                else:
                    sd1.add_paragraph(input.paragraphs[x].text, input.paragraphs[x].style)

            # sd1 = style_match(input.paragraphs[start+1:end], sd1)
            replace[p] = highlight_TBDs(sd1)

    
    if not replace['pm'] is None:
        pm_sig = find_signature(replace['pm'])
        pm_image = InlineImage(template, image_descriptor=pm_sig, width= Mm(30), height= Mm(15))
        replace['pm_signature'] = pm_image
    if not replace['pc'] is None:
        pc_sig = find_signature(replace['pc'])
        pc_image = InlineImage(template, image_descriptor=pc_sig, width= Mm(30), height= Mm(15))
        replace['pc_signature'] = pc_image



    path, fn = os.path.split(input_path)
    name, ext = os.path.splitext(fn)
    new_fn = name + '_REPORT' + ext


    if save_path is None:
        save_path = os.path.join(path, new_fn)
    else:
        save_path = os.path.join(save_path, new_fn)

    template.render(replace)

    if save:
        template.save(save_path)

    return template, save_path

    
        

def scrape_test(input_path, template_path, save_path = None, save = True):
    input = Document(input_path)
    template = DocxTemplate(template_path)

    # tab_dict = extract_tabs(input_path)

    # Pull data from proposal (input)
    data_ind = document_TOC(input) 
    d1 = scrape_project_data(0, end, input)
    # title_page =[(ind, p.text) for ind, p in enumerate(input.paragraphs) if ind < data_ind['TABLE OF CONTENTS'] and p.text.strip() != '' ]    
    # table_loc, table = document_tables(input)

    # Pull placeholders from template
    place = template.get_undeclared_template_variables()
    replace = {k : None for k in place}



    # Identify key points, update replace dict
    # replace['title'] = title_page[0][1]
    # replace['study_num'] = [ x[1].split(': ')[-1] for x in title_page if 'Project Quotation' in x[1]][0]
    replace['study_num'] = input.sections[0].header.paragraphs[1].text.split('\t')[0]
    # replace['client'] = [x[1].split('Prepared for ')[-1] for x in title_page if 'Prepared for' in x[1]][0]
    # replace['t'] = date.today().strftime("%B %d, %Y")
    replace['description'] = input.sections[0].header.paragraphs[0].text.split('\t')[0]

    # Replace '&' in strings to prevent jinga2 errors
    for k,v in replace.items():
        if v is not None and '&' in v:
            replace[k] = v.replace("&","and")

    # styles = input.styles 
    # paragraph_styles = [s for s in styles if s.type == WD_STYLE_TYPE.PARAGRAPH ]
    # sd1 = template.new_subdoc()
    # for style in paragraph_styles:
    #     sd1.add_paragraph('style', style)
    # replace['styles'] = sd1

    # For section replacement, iterate through and pull paragraphs from proposal
    for p in place:
        if p in data_ind.keys():
            
            start = data_ind[p][0]
            end = data_ind[p][1]
            # print(p+": "+str(start)+":"+str(end))
            sd1 = template.new_subdoc()
            
            for x in range(start+1, end, 1):
                
                if 'Bullet' in input.paragraphs[x].style.name:
                    para = sd1.add_paragraph(input.paragraphs[x].text, input.paragraphs[x].style)

                    # if x in tab_dict.keys():
                    #     lev = tab_dict[x]
                    # else:
                    #     lev = None
                    lev = None
                    # if x == start+1:
                    list_number(sd1, para, prev=input.paragraphs[x], level=lev, num=False)
                    # else:
                        # list_number(sd1, para, prev=input.paragraphs[x-1], level=lev, num=False)
                else:
                    sd1.add_paragraph(input.paragraphs[x].text, input.paragraphs[x].style)


            # print(start)
            # print(end)
            # sd1 = style_match(input.paragraphs[int(start+1) :int(end)], sd1)
            replace[p] = sd1
    
    path, fn = os.path.split(input_path)
    name, ext = os.path.splitext(fn)
    new_fn = name + '_REPORT' + ext


    if save_path is None:
        save_path = os.path.join(path, new_fn)
    else:
        save_path = os.path.join(save_path, new_fn)


    template.render(replace)

    if save:
        template.save(save_path)


